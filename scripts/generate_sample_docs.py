"""Generate 15 service desk sample documents (8 Word + 7 PDF) for MkDocs learning."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    ListFlowable,
    ListItem,
)

ROOT = Path(__file__).resolve().parents[1]
WORD_DIR = ROOT / "sample-data" / "word"
PDF_DIR = ROOT / "sample-data" / "pdf"
WORD_DIR.mkdir(parents=True, exist_ok=True)
PDF_DIR.mkdir(parents=True, exist_ok=True)

WORD_DOCS = [
    {
        "filename": "01-service-desk-overview.docx",
        "title": "Service Desk Overview",
        "meta": "IT Service Management | Audience: All staff | Version 1.0",
        "sections": [
            (
                "Purpose",
                [
                    "This document introduces Contoso IT Service Desk: how to request help, what we support, and how work is tracked.",
                    "Use this as the first page of your MkDocs service desk site (home or getting-started section).",
                ],
                [
                    "Single point of contact for IT issues and requests",
                    "Tickets logged in the service management tool",
                    "Self-service knowledge base encouraged before opening a ticket",
                ],
                None,
            ),
            (
                "How to contact us",
                [
                    "Choose the channel that matches urgency. Critical outages should use phone first, then open a ticket."
                ],
                [
                    "Portal: https://support.contoso.example (preferred)",
                    "Email: servicedesk@contoso.example",
                    "Phone: +1 (555) 010-2000 (business hours)",
                    "Chat: available in the employee portal 08:00-17:00 local time",
                ],
                None,
            ),
            (
                "What we support",
                [
                    "Supported services include accounts, email, endpoint devices, VPN, collaboration tools, and standard line-of-business applications."
                ],
                [
                    "Windows and managed macOS endpoints",
                    "Microsoft 365 (Outlook, Teams, SharePoint, OneDrive)",
                    "Corporate VPN and zero-trust access",
                    "Standard printers and conference room AV (via facilities for hardware faults)",
                ],
                None,
            ),
            (
                "Business hours",
                [
                    "Standard support is Monday-Friday, 08:00-18:00 local time, excluding company holidays. After-hours coverage is limited to Severity 1 incidents."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "02-password-reset-procedure.docx",
        "title": "Password Reset Procedure",
        "meta": "SOP-AUTH-001 | Audience: End users and Tier 1 | Version 1.2",
        "sections": [
            (
                "Overview",
                [
                    "Users should reset passwords via self-service whenever possible. Agents follow this procedure when self-service fails or the account is locked."
                ],
                [],
                None,
            ),
            (
                "Self-service steps",
                ["Direct users to complete these steps before escalating to an agent."],
                [
                    "Go to https://password.contoso.example",
                    "Enter work email address and complete MFA challenge",
                    "Choose a new password that meets complexity rules (14+ characters, mixed case, number, symbol)",
                    "Sign out of all sessions and sign back in with the new password",
                ],
                None,
            ),
            (
                "Agent procedure",
                ["If self-service is unavailable, verify identity before resetting."],
                [
                    "Open the ticket and confirm employee ID, manager, and last four of corporate mobile (or HR-approved alternate)",
                    "In the identity admin console, select Reset password and force change at next sign-in",
                    "Clear temporary lockouts if present",
                    "Provide the temporary password via a secure channel (never email permanent secrets)",
                    "Advise user to complete MFA re-registration if prompts fail after reset",
                    "Document steps taken and close or resolve the ticket",
                ],
                None,
            ),
            (
                "Common issues",
                [
                    "Password accepted but apps still fail: clear cached credentials, restart the device, and re-authenticate VPN/MFA apps."
                ],
                [
                    "Account locked: wait 15 minutes or clear lock in admin console",
                    "MFA device lost: escalate to Identity team with manager approval",
                    "Shared/service accounts: do not reset without owner approval",
                ],
                None,
            ),
        ],
    },
    {
        "filename": "03-vpn-access-guide.docx",
        "title": "VPN Access Guide",
        "meta": "KB-NET-014 | Audience: Remote workers | Version 2.0",
        "sections": [
            (
                "When to use VPN",
                [
                    "Use the corporate VPN only when you need access to internal systems that are not published through the secure web gateway (legacy file shares, certain finance apps, lab networks)."
                ],
                [],
                None,
            ),
            (
                "Install and connect",
                [],
                [
                    "Install Contoso Secure Access from Company Portal / Managed Software Center",
                    "Sign in with your work account (SSO)",
                    "Approve the MFA prompt on your authenticator app",
                    "Confirm status shows Connected and your internal IP is assigned",
                ],
                None,
            ),
            (
                "Troubleshooting",
                [
                    "If connection fails, collect the client logs from Help > Export logs before contacting the service desk."
                ],
                [
                    "Home Wi-Fi blocking UDP: try alternate port profile or wired connection",
                    "Certificate errors: restart device and ensure date/time is automatic",
                    "Split tunnel: internet should still work; only corporate routes go through VPN",
                ],
                None,
            ),
            (
                "Security rules",
                [
                    "Do not share VPN profiles. Personal devices require enrollment in mobile device management before VPN is allowed."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "04-incident-severity-levels.docx",
        "title": "Incident Severity Levels",
        "meta": "POL-IM-002 | Audience: Service desk and IT ops | Version 1.4",
        "sections": [
            (
                "Purpose",
                [
                    "Severity drives response time, communication cadence, and escalation. Agents must set severity based on business impact, not user preference alone."
                ],
                [],
                [
                    ["Severity", "Definition", "Initial response", "Example"],
                    [
                        "Sev 1",
                        "Critical service down for many users or revenue impact",
                        "15 minutes",
                        "Email outage company-wide",
                    ],
                    [
                        "Sev 2",
                        "Major function degraded; workaround limited",
                        "1 hour",
                        "VPN down for one region",
                    ],
                    [
                        "Sev 3",
                        "Single user or small group; business can continue",
                        "4 hours",
                        "One laptop cannot print",
                    ],
                    [
                        "Sev 4",
                        "Low impact request or cosmetic issue",
                        "1 business day",
                        "Desktop shortcut missing",
                    ],
                ],
            ),
            (
                "How to choose severity",
                [],
                [
                    "Count of users affected",
                    "Whether a viable workaround exists",
                    "Regulatory, safety, or security implications",
                    "Whether the issue is in a production vs. test environment",
                ],
                None,
            ),
            (
                "Communication",
                [
                    "Sev 1 and Sev 2 require status updates on the major-incident channel until resolved or mitigated."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "05-software-request-process.docx",
        "title": "Software Installation Request Process",
        "meta": "SOP-REQ-008 | Audience: Employees and approvers | Version 1.1",
        "sections": [
            (
                "Policy summary",
                [
                    "Only approved software may be installed on managed devices. Unlisted titles require security and licensing review."
                ],
                [],
                None,
            ),
            (
                "Request flow",
                [],
                [
                    "Search the software catalog for an existing package",
                    "If listed: click Request install (auto-approves for free tier tools)",
                    "If not listed: open a Service Request with business justification, cost center, and data classification",
                    "Manager approves; Security reviews if the tool handles regulated data",
                    "Packaging team publishes to Company Portal when approved",
                ],
                None,
            ),
            (
                "SLA targets",
                [
                    "Catalog apps: same day. New packaging: 5-10 business days depending on complexity."
                ],
                [],
                [
                    ["Request type", "Target"],
                    ["Catalog install", "8 business hours"],
                    ["License add-on", "2 business days"],
                    ["New package (standard)", "5 business days"],
                    ["New package (security review)", "10 business days"],
                ],
            ),
        ],
    },
    {
        "filename": "06-new-hire-it-checklist.docx",
        "title": "New Employee IT Onboarding Checklist",
        "meta": "CHK-HRIT-003 | Audience: Service desk / HR partners | Version 3.0",
        "sections": [
            (
                "Before day one",
                [],
                [
                    "Create identity account from HR ticket (legal name, department, manager, start date)",
                    "Assign license bundle based on role template",
                    "Stage laptop image and ship or stage for pickup",
                    "Add to required security groups (department + VPN if remote)",
                    "Schedule welcome email with first-login instructions",
                ],
                None,
            ),
            (
                "Day one",
                [],
                [
                    "Verify MFA enrollment completed",
                    "Confirm email, Teams, and OneDrive sign-in",
                    "Map standard printers if site-based",
                    "Walk through password self-service registration",
                    "Share link to service desk knowledge base home",
                ],
                None,
            ),
            (
                "Within first week",
                [],
                [
                    "Confirm access to role-specific applications",
                    "Close onboarding ticket only after user confirmation",
                    "Tag ticket with site and department for reporting",
                ],
                None,
            ),
        ],
    },
    {
        "filename": "07-email-troubleshooting.docx",
        "title": "Email Troubleshooting Guide",
        "meta": "KB-M365-021 | Audience: Tier 1 agents | Version 1.6",
        "sections": [
            (
                "Quick triage",
                [
                    "Determine whether the issue is send, receive, calendar, or mobile-only before changing settings."
                ],
                [
                    "Can the user sign in at outlook.office.com?",
                    "Does the issue affect one device or all devices?",
                    "Any recent password or MFA changes?",
                    "Is the mailbox full (quota warning)?",
                ],
                None,
            ),
            (
                "Desktop Outlook",
                [],
                [
                    "Repair Office apps from Settings > Apps",
                    "Create a new Outlook profile if the profile is corrupt",
                    "Clear credential manager entries for MicrosoftOffice*",
                    "Check rules and Focused Inbox for missing mail",
                ],
                None,
            ),
            (
                "Mobile mail",
                [],
                [
                    "Remove and re-add the work account in the Outlook mobile app",
                    "Confirm device is compliant in Company Portal",
                    "Do not recommend native iOS/Android mail for corporate accounts",
                ],
                None,
            ),
            (
                "Escalate when",
                [
                    "Mailbox corruption, litigation hold issues, transport rule failures, or hybrid mail flow problems - assign to Messaging team with logs attached."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "08-remote-desktop-sop.docx",
        "title": "Remote Desktop Access SOP",
        "meta": "SOP-EUC-011 | Audience: Support agents | Version 1.0",
        "sections": [
            (
                "Policy",
                [
                    "Interactive remote control requires user consent except for break-glass admin scenarios approved by the duty manager."
                ],
                [],
                None,
            ),
            (
                "Starting a remote session",
                [],
                [
                    "Open the remote support tool from the agent console",
                    "Enter the ticket number (required for audit)",
                    "Send the session invitation to the user",
                    "Wait for Accept on the user device",
                    "Announce actions before making system changes",
                    "End session and note outcome in the ticket",
                ],
                None,
            ),
            (
                "Prohibited actions",
                [],
                [
                    "Do not disable security tools without Security approval",
                    "Do not copy personal files off the device",
                    "Do not share your admin credentials with the user",
                ],
                None,
            ),
        ],
    },
]

PDF_DOCS = [
    {
        "filename": "09-sla-overview.pdf",
        "title": "Service Level Agreement Overview",
        "meta": "POL-SLA-001 | Version 2.1",
        "sections": [
            (
                "Scope",
                [
                    "This SLA covers Contoso IT services delivered through the Service Desk for full-time employees and long-term contractors."
                ],
                [
                    "Incident response and resolution targets by severity",
                    "Service request fulfillment targets",
                    "Planned maintenance windows",
                ],
                None,
            ),
            (
                "Availability targets",
                [
                    "Core collaboration services target 99.9% monthly availability excluding announced maintenance."
                ],
                [],
                [
                    ["Service", "Target availability"],
                    ["Email / Teams", "99.9%"],
                    ["VPN / remote access", "99.5%"],
                    ["Service desk portal", "99.5%"],
                    ["Identity / SSO", "99.95%"],
                ],
            ),
            (
                "Exclusions",
                [
                    "Force majeure, customer-owned networks, and unsanctioned software are excluded from SLA credits."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "10-escalation-matrix.pdf",
        "title": "Escalation Matrix",
        "meta": "REF-IM-010 | Version 1.3",
        "sections": [
            (
                "When to escalate",
                [
                    "Escalate when the SLA is at risk, the issue exceeds Tier 1 skills, or a major incident is declared."
                ],
                [],
                [
                    ["Level", "Owner", "Trigger"],
                    ["Tier 1", "Service Desk", "Initial triage and known fixes"],
                    ["Tier 2", "Platform teams", "Complex config, multi-user impact"],
                    ["Tier 3", "Engineering / vendor", "Code defects, infrastructure faults"],
                    ["Major Incident", "Duty manager", "Sev 1 or widespread Sev 2"],
                ],
            ),
            (
                "Contacts (sample)",
                [],
                [
                    "Duty manager on-call: via PagerDuty schedule IT-Duty",
                    "Security CSIRT: security@contoso.example",
                    "Vendor bridge: open only after Tier 3 acknowledges",
                ],
                None,
            ),
            (
                "Documentation required",
                [
                    "Every escalation must include timeline, impact statement, reproduction steps, and current workaround."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "11-kb-authoring-guidelines.pdf",
        "title": "Knowledge Base Authoring Guidelines",
        "meta": "STD-KB-004 | Audience: Authors | Version 1.0",
        "sections": [
            (
                "Why structure matters",
                [
                    "MkDocs works best when articles use consistent headings, short procedures, and clear titles. These source files are practice content to convert into Markdown pages."
                ],
                [],
                None,
            ),
            (
                "Article template",
                [],
                [
                    "Title: task-oriented (e.g., Reset your VPN client)",
                    "Summary: 1-2 sentences",
                    "Before you begin: prerequisites",
                    "Steps: numbered, one action each",
                    "Result: what success looks like",
                    "Related articles: links",
                ],
                None,
            ),
            (
                "Style rules",
                [],
                [
                    "Use active voice and second person (you)",
                    "Prefer screenshots only when UI is non-obvious",
                    "Mark severity of risk with callouts (warning/note)",
                    "Review quarterly or after major product changes",
                ],
                None,
            ),
            (
                "MkDocs learning tip",
                [
                    "When you convert this PDF to Markdown, map each section heading to ## or ### and turn bullet lists into Markdown lists. Place the page under docs/service-desk/ in your MkDocs project."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "12-hardware-request-guide.pdf",
        "title": "Hardware Asset Request Guide",
        "meta": "SOP-AM-006 | Version 1.2",
        "sections": [
            (
                "Eligible hardware",
                [
                    "Standard laptop, monitor, headset, docking station, and keyboard/mouse kits are ordered from the asset catalog."
                ],
                [],
                [
                    ["Item", "Refresh cycle", "Notes"],
                    ["Laptop", "3-4 years", "Role-based models only"],
                    ["Monitor", "5 years", "Max two per desk by default"],
                    ["Headset", "As needed", "Noise-cancelling standard"],
                ],
            ),
            (
                "How to request",
                [],
                [
                    "Open Service Request > Hardware",
                    "Select catalog item and ship-to address",
                    "Manager approval required for non-standard items",
                    "Asset tag is recorded on delivery; return old kit if refresh",
                ],
                None,
            ),
            (
                "Lost or stolen devices",
                [
                    "Report immediately to Service Desk and Security. Do not delay for a replacement request."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "13-change-management-basics.pdf",
        "title": "Change Management Basics",
        "meta": "POL-CHG-001 | Version 1.5",
        "sections": [
            (
                "Goal",
                [
                    "Changes to production IT services are planned, reviewed, and communicated to reduce unplanned downtime."
                ],
                [],
                None,
            ),
            (
                "Change types",
                [],
                [],
                [
                    ["Type", "Approval", "Lead time"],
                    ["Standard", "Pre-approved", "As scheduled"],
                    ["Normal", "CAB / manager", "3 business days"],
                    ["Emergency", "Duty manager", "As needed"],
                ],
            ),
            (
                "Service desk role",
                [],
                [
                    "Communicate maintenance windows to users",
                    "Monitor incidents during change implementation",
                    "Link related incidents to the change record",
                ],
                None,
            ),
        ],
    },
    {
        "filename": "14-security-incident-quickref.pdf",
        "title": "Security Incident Quick Reference",
        "meta": "REF-SEC-002 | Version 1.1 | Confidential - internal",
        "sections": [
            (
                "If you suspect a security incident",
                [
                    "Examples: phishing that was clicked, ransomware symptoms, lost device with data, unexpected admin prompts, data emailed externally by mistake."
                ],
                [
                    "Do not power off a potentially compromised device unless instructed (preserve evidence) - disconnect network if malware is active",
                    "Contact Service Desk and mark ticket Security",
                    "Notify Security CSIRT for confirmed or high-confidence events",
                    "Do not delete emails or files that may be evidence",
                ],
                None,
            ),
            (
                "Phishing reports",
                [],
                [
                    "Use Report phishing button in Outlook",
                    "Forward as attachment only if button missing",
                    "Never enter credentials on linked pages from suspicious mail",
                ],
                None,
            ),
            (
                "Service desk first response",
                [
                    "Isolate account if credential theft is likely: reset password, revoke sessions, review MFA methods. Escalate to CSIRT with timeline."
                ],
                [],
                None,
            ),
        ],
    },
    {
        "filename": "15-printer-peripheral-support.pdf",
        "title": "Printer and Peripheral Support",
        "meta": "KB-EUC-033 | Version 1.0",
        "sections": [
            (
                "Supported devices",
                [
                    "Corporate networked printers, standard USB headsets, webcams, and docking stations listed in the hardware catalog."
                ],
                [],
                None,
            ),
            (
                "Printer issues checklist",
                [],
                [
                    "Confirm user is on corporate network or VPN if required",
                    "Reinstall printer from print server share",
                    "Clear stuck jobs from the print queue",
                    "Check toner/paper and physical error lights",
                    "Escalate facilities for paper jams requiring hardware service",
                ],
                None,
            ),
            (
                "Docking and monitors",
                [
                    "Update dock firmware via Company Portal. Try a different USB-C/Thunderbolt cable before replacing hardware."
                ],
                [],
                None,
            ),
            (
                "MkDocs practice idea",
                [
                    "Split this article into two Markdown pages later: printers.md and docks-monitors.md, and link them from a peripherals index page."
                ],
                [],
                None,
            ),
        ],
    },
]


def add_word_doc(spec):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(spec["title"], level=0)
    meta = doc.add_paragraph(spec["meta"])
    meta.runs[0].italic = True
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note = doc.add_paragraph()
    run = note.add_run(
        "Sample training content for learning MkDocs: convert this document into Markdown under docs/."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8A)

    for heading, paragraphs, bullets, table in spec["sections"]:
        doc.add_heading(heading, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
        if table:
            t = doc.add_table(rows=len(table), cols=len(table[0]))
            t.style = "Table Grid"
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    t.rows[i].cells[j].text = cell
                    if i == 0:
                        for paragraph in t.rows[i].cells[j].paragraphs:
                            for r in paragraph.runs:
                                r.bold = True

    path = WORD_DIR / spec["filename"]
    doc.save(path)
    return path


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="DocTitle",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=HexColor("#1a365d"),
            spaceAfter=8,
            leading=22,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            parent=styles["Normal"],
            fontSize=9,
            textColor=HexColor("#555555"),
            spaceAfter=12,
            fontName="Helvetica-Oblique",
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=8
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2",
            parent=styles["Heading2"],
            fontSize=13,
            textColor=HexColor("#2c5282"),
            spaceBefore=14,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Tip",
            parent=styles["Normal"],
            fontSize=9,
            textColor=HexColor("#1a568a"),
            spaceAfter=10,
            leading=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletBody", parent=styles["Normal"], fontSize=10, leading=13
        )
    )
    return styles


def add_pdf_doc(spec, styles):
    path = PDF_DIR / spec["filename"]
    pdf = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    story = []
    story.append(Paragraph(spec["title"], styles["DocTitle"]))
    story.append(Paragraph(spec["meta"], styles["Meta"]))
    story.append(
        Paragraph(
            "Sample training content for learning MkDocs: convert this PDF into Markdown under docs/.",
            styles["Tip"],
        )
    )
    story.append(Spacer(1, 6))

    for heading, paragraphs, bullets, table in spec["sections"]:
        story.append(Paragraph(heading, styles["H2"]))
        for p in paragraphs:
            story.append(Paragraph(p, styles["Body"]))
        if bullets:
            items = [
                ListItem(Paragraph(b, styles["BulletBody"]), leftIndent=12)
                for b in bullets
            ]
            story.append(ListFlowable(items, bulletType="bullet", start="•"))
            story.append(Spacer(1, 6))
        if table:
            data = [
                [Paragraph(str(c), styles["BulletBody"]) for c in row] for row in table
            ]
            ncols = len(table[0])
            col_w = (6.5 * inch) / ncols
            t = Table(data, hAlign="LEFT", colWidths=[col_w] * ncols)
            t.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#2c5282")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("BACKGROUND", (0, 1), (-1, -1), HexColor("#f7fafc")),
                        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cbd5e0")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 8))

    pdf.build(story)
    return path


def main():
    styles = build_pdf_styles()
    created = []
    for spec in WORD_DOCS:
        created.append(add_word_doc(spec))
    for spec in PDF_DOCS:
        created.append(add_pdf_doc(spec, styles))
    print(f"Created {len(created)} files:")
    for p in created:
        print(f"  {p.relative_to(ROOT)} ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
